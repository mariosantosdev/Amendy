from os import system, name # Inserir codigo no terminal & verificar o SO
from docx import Document
from docx.shared import Inches

def systemDetect(): # Verificar SO
    if(name == 'nt'):
        return('win')
    else:
        return('other')

def clearTerminal(): # Limpar terminal
    if(systemDetect() == 'win'):
        system('cls')
    else:
        system('clear')

def createFile(content, theme): # Cria o arquivo do word
    doc = Document()

    doc.add_heading(theme)
    doc.add_paragraph(content)

    saveName = theme.replace(' ', '-')
    return doc.save(f'./Projects/{saveName}.docx')
